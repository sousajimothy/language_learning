"""
pages/5_Report.py — Worst-performing items report.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
import plotly.graph_objects as go
import streamlit as st

from german_pipeline import storage
from ui_utils import (
    build_plotly_layout, cutoff_iso, fmt_rate, fmt_ts,
    get_db_path, get_plotly_colors, list_sources, open_db,
)

# ---------------------------------------------------------------------------
# CSS
# ---------------------------------------------------------------------------

st.markdown("""
<style>
.report-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 1.1rem;
}
.report-title {
    font-size: 1.5rem;
    font-weight: 800;
    color: var(--text-color);
    letter-spacing: -0.01em;
}
.filter-pill {
    font-size: 0.65rem;
    font-weight: 700;
    letter-spacing: 0.10em;
    text-transform: uppercase;
    padding: 0.22rem 0.65rem;
    border-radius: 4px;
    background: rgba(49,130,206,0.10);
    color: #3182CE;
    border: 1px solid rgba(49,130,206,0.20);
}
.section-lbl {
    font-size: 0.64rem;
    font-weight: 700;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: color-mix(in srgb, var(--text-color) 28%, transparent);
    margin: 1.5rem 0 0.65rem;
    display: flex;
    align-items: center;
    gap: 0.55rem;
}
.section-lbl::after {
    content: '';
    flex: 1;
    height: 1px;
    background: color-mix(in srgb, var(--text-color) 6%, transparent);
}
.stat-card {
    background: var(--secondary-background-color);
    border: 1.5px solid rgba(128,128,128,0.40);
    border-radius: 10px;
    padding: 0.85rem 1rem;
    text-align: center;
}
.stat-card .sc-num {
    font-size: 1.8rem;
    font-weight: 800;
    color: var(--text-color);
    line-height: 1;
}
.stat-card .sc-lbl {
    font-size: 0.62rem;
    font-weight: 700;
    letter-spacing: 0.09em;
    text-transform: uppercase;
    color: color-mix(in srgb, var(--text-color) 30%, transparent);
    margin-top: 0.28rem;
}
.stat-card .sc-sub {
    font-size: 0.70rem;
    color: color-mix(in srgb, var(--text-color) 22%, transparent);
    margin-top: 0.12rem;
}
.export-row {
    display: flex;
    justify-content: flex-end;
    margin-top: 0.4rem;
}
/* ── Download cards ─────────────────────────────────────────────── */
.dl-card {
    background: color-mix(in srgb, var(--text-color) 4%, transparent);
    border: 1px solid color-mix(in srgb, var(--text-color) 10%, transparent);
    border-radius: 8px;
    padding: 0.9rem 1rem 0.6rem;
    margin-bottom: 0.5rem;
}
.dl-card .dl-title {
    font-weight: 600;
    font-size: 0.9rem;
    color: color-mix(in srgb, var(--text-color) 85%, transparent);
    margin-bottom: 0.15rem;
}
.dl-card .dl-desc {
    font-size: 0.73rem;
    color: color-mix(in srgb, var(--text-color) 38%, transparent);
    margin-bottom: 0.65rem;
    line-height: 1.4;
}
</style>
""", unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Multi-format export helper
# ---------------------------------------------------------------------------

def _build_exports(
    df: pd.DataFrame,
    title: str,
    file_stem: str,
    pct_cols: list[str] | None = None,
    key_prefix: str = "exp",
) -> None:
    """Render a 3×2 download-card grid for *df* in 6 formats.

    *pct_cols* lists column names whose values are 0–1 floats that should
    be formatted as percentages in the human-readable exports.
    """
    pct_cols = pct_cols or []
    _export_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    _cols = list(df.columns)
    _n = len(df)

    # ── Percentage-formatted copy for human-readable exports ────────
    hdf = df.copy()
    for c in pct_cols:
        if c in hdf.columns:
            hdf[c] = hdf[c].apply(lambda v: f"{v:.0%}" if pd.notna(v) else "—")

    # ── CSV ─────────────────────────────────────────────────────────
    csv_buf = io.StringIO()
    df.to_csv(csv_buf, index=False, encoding="utf-8")
    csv_bytes = ("\ufeff" + csv_buf.getvalue()).encode("utf-8")

    # ── TSV ─────────────────────────────────────────────────────────
    tsv_buf = io.StringIO()
    df.to_csv(tsv_buf, sep="\t", index=False, encoding="utf-8")
    tsv_bytes = tsv_buf.getvalue().encode("utf-8")

    # ── XLSX ────────────────────────────────────────────────────────
    xlsx_buf = io.BytesIO()
    df.to_excel(xlsx_buf, index=False, engine="openpyxl")
    xlsx_bytes = xlsx_buf.getvalue()

    # ── Markdown (GFM pipe table) ──────────────────────────────────
    def _md_cell(v: str) -> str:
        return v.replace("|", "\\|").replace("\n", " ")

    md_parts = [
        f"# {title}\n",
        f"_{_n} item{'s' if _n != 1 else ''} · Exported {_export_date}_\n",
        "",
        "| " + " | ".join(_md_cell(c) for c in _cols) + " |",
        "| " + " | ".join("---" for _ in _cols) + " |",
    ]
    for _, row in hdf.iterrows():
        md_parts.append(
            "| " + " | ".join(_md_cell(str(row.get(c, "") or "")) for c in _cols) + " |"
        )
    md_bytes = "\n".join(md_parts).encode("utf-8")

    # ── DOCX ────────────────────────────────────────────────────────
    import docx as _docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Ox

    _DH_BG   = "3B5998"
    _DH_TEXT  = RGBColor(0xFF, 0xFF, 0xFF)
    _DT_CLR   = RGBColor(0x2D, 0x3A, 0x4A)
    _DS_CLR   = RGBColor(0x8B, 0x95, 0xA5)
    _DB_CLR   = RGBColor(0x2D, 0x3A, 0x4A)

    def _d_shade(cell, hx):
        shd = _Ox("w:shd"); shd.set(_qn("w:fill"), hx); shd.set(_qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)

    def _d_tbl_borders(tbl):
        tblPr = tbl._tbl.tblPr
        def _b(tag, color, sz=4, val="single"):
            b = _Ox(f"w:{tag}"); b.set(_qn("w:val"), val); b.set(_qn("w:sz"), str(sz)); b.set(_qn("w:color"), color)
            return b
        bdr = _Ox("w:tblBorders")
        for tag, clr, sz, val in [("top","CFD8DC",4,"single"),("left","CFD8DC",2,"single"),("bottom","CFD8DC",4,"single"),
                                   ("right","CFD8DC",2,"single"),("insideH","E8ECF0",2,"single"),("insideV","none",0,"none")]:
            bdr.append(_b(tag, clr, sz, val))
        tblPr.append(bdr)

    def _d_cell_margins(tbl, top=100, start=140, bottom=100, end=140):
        tblPr = tbl._tbl.tblPr
        mar = _Ox("w:tblCellMar")
        for side, val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
            n = _Ox(f"w:{side}"); n.set(_qn("w:w"), str(val)); n.set(_qn("w:type"), "dxa")
            mar.append(n)
        tblPr.append(mar)

    def _d_para_rule(para, color="B0BEC5"):
        pPr = para._p.get_or_add_pPr()
        pBdr = _Ox("w:pBdr"); b = _Ox("w:bottom")
        b.set(_qn("w:val"), "single"); b.set(_qn("w:sz"), "6"); b.set(_qn("w:color"), color); b.set(_qn("w:space"), "8")
        pBdr.append(b); pPr.append(pBdr)

    def _d_repeat_header(tbl):
        trPr = tbl.rows[0]._tr.get_or_add_trPr(); trPr.append(_Ox("w:tblHeader"))

    _doc = _docx.Document()
    _doc.core_properties.title = title
    _section = _doc.sections[0]
    _section.page_width  = Cm(29.7); _section.page_height = Cm(21.0)
    _section.left_margin = Cm(2.5);  _section.right_margin = Cm(2.5)
    _section.top_margin  = Cm(2.0);  _section.bottom_margin = Cm(1.8)

    # Title
    _tp = _doc.add_paragraph()
    _tp.paragraph_format.space_before = Pt(0); _tp.paragraph_format.space_after = Pt(4)
    _tr = _tp.add_run(title); _tr.bold = True; _tr.font.size = Pt(20); _tr.font.color.rgb = _DT_CLR
    _d_para_rule(_tp)

    # Subtitle — tab-stop layout
    _sp = _doc.add_paragraph()
    _sp.paragraph_format.space_before = Pt(8); _sp.paragraph_format.space_after = Pt(20)
    _usable_w = _section.page_width - _section.left_margin - _section.right_margin
    _sp.paragraph_format.tab_stops.add_tab_stop(_usable_w, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _sr_l = _sp.add_run(f"{_n} item{'s' if _n != 1 else ''}"); _sr_l.font.size = Pt(9.5); _sr_l.font.color.rgb = _DS_CLR
    _sp.add_run("\t")
    _sr_r = _sp.add_run(f"Exported {_export_date}"); _sr_r.font.size = Pt(9.5); _sr_r.font.color.rgb = _DS_CLR

    # Table — distribute widths evenly across usable space
    _n_cols = len(_cols)
    _per_col = Cm(24.7 / _n_cols)
    _tbl = _doc.add_table(rows=1, cols=_n_cols)
    _tbl.style = "Table Grid"; _tbl.autofit = False; _tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _d_cell_margins(_tbl); _d_tbl_borders(_tbl)
    for i in range(_n_cols):
        _tbl.columns[i].width = _per_col

    _hdr_cells = _tbl.rows[0].cells
    for i, col_name in enumerate(_cols):
        _hdr_cells[i].text = col_name
        _hdr_cells[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        _d_shade(_hdr_cells[i], _DH_BG)
        _run = _hdr_cells[i].paragraphs[0].runs[0]
        _run.bold = True; _run.font.size = Pt(9); _run.font.color.rgb = _DH_TEXT
        _hdr_cells[i].paragraphs[0].paragraph_format.space_before = Pt(3)
        _hdr_cells[i].paragraphs[0].paragraph_format.space_after  = Pt(3)
    _d_repeat_header(_tbl)

    for _row_idx, (_, _data_row) in enumerate(hdf.iterrows()):
        _row_cells = _tbl.add_row().cells
        for i, col_name in enumerate(_cols):
            _row_cells[i].text = str(_data_row.get(col_name, "") or "")
            if _row_idx % 2 == 1:
                _d_shade(_row_cells[i], "F5F7FA")
            _run = _row_cells[i].paragraphs[0].runs[0]
            _run.font.size = Pt(9.5); _run.font.color.rgb = _DB_CLR
            _row_cells[i].paragraphs[0].paragraph_format.space_before = Pt(1)
            _row_cells[i].paragraphs[0].paragraph_format.space_after  = Pt(1)

    docx_buf = io.BytesIO(); _doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    # ── PDF ─────────────────────────────────────────────────────────
    from fpdf import FPDF

    _FONTS_DIR    = Path(__file__).resolve().parent.parent / "assets" / "fonts"
    _FONT_REGULAR = _FONTS_DIR / "DejaVuSans.ttf"
    _FONT_BOLD    = _FONTS_DIR / "DejaVuSans-Bold.ttf"
    if not _FONT_REGULAR.exists() or not _FONT_BOLD.exists():
        st.error("PDF export requires bundled fonts in `assets/fonts/`.")
        st.stop()

    _PH  = (59, 89, 152);  _PHD = (42, 63, 110)
    _PW  = (255, 255, 255); _PT  = (45, 58, 74)
    _PS  = (139, 149, 165); _PA  = (245, 247, 250)
    _PBR = (207, 216, 220); _PBL = (232, 236, 240)

    _total_w   = 257  # landscape A4 minus 20 mm margins each side
    _col_w_pdf = [round(_total_w / _n_cols)] * _n_cols
    # distribute rounding remainder
    _col_w_pdf[-1] = _total_w - sum(_col_w_pdf[:-1])
    _line_h = 7; _hdr_h = 10

    def _pdf_col_headers(pdf, cols, widths, h):
        pdf.set_font("DejaVu", style="B", size=9)
        pdf.set_fill_color(*_PH); pdf.set_text_color(*_PW)
        x0, y0 = pdf.l_margin, pdf.get_y()
        pdf.rect(x0, y0, sum(widths), h, "F")
        for col_name, w in zip(cols, widths):
            pdf.multi_cell(w, h, f"  {col_name}", border=0, fill=False,
                           new_x="RIGHT", new_y="TOP", max_line_height=h, align="L")
        pdf.set_y(y0 + h)
        pdf.set_draw_color(*_PHD); pdf.set_line_width(0.6)
        pdf.line(x0, pdf.get_y(), x0 + sum(widths), pdf.get_y())
        pdf.set_draw_color(*_PBL); pdf.set_line_width(0.15)
        pdf.set_text_color(*_PT)

    class _ReportPDF(FPDF):
        def footer(self):
            self.set_y(-14)
            self.set_draw_color(*_PBR); self.set_line_width(0.3)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.set_y(-12); self.set_font("DejaVu", size=7.5); self.set_text_color(*_PS)
            self.cell(0, 5, f"Page {self.page_no()} / {{nb}}", align="R")
            self.set_text_color(*_PT)

    _pdf = _ReportPDF(orientation="L", unit="mm", format="A4")
    _pdf.add_font("DejaVu", style="", fname=str(_FONT_REGULAR))
    _pdf.add_font("DejaVu", style="B", fname=str(_FONT_BOLD))
    _pdf.alias_nb_pages()
    _pdf.set_margins(left=20, top=18, right=20)
    _pdf.set_auto_page_break(auto=False)
    _pdf.add_page()

    # Title block
    _pdf.set_font("DejaVu", style="B", size=20); _pdf.set_text_color(*_PT)
    _pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
    _pdf.set_draw_color(*_PBR); _pdf.set_line_width(0.5)
    _pdf.line(_pdf.l_margin, _pdf.get_y() + 1, _pdf.w - _pdf.r_margin, _pdf.get_y() + 1)
    _pdf.set_line_width(0.15); _pdf.ln(5)

    _pdf.set_font("DejaVu", size=10); _pdf.set_text_color(*_PS)
    _pdf.cell(0, 6, f"{_n} item{'s' if _n != 1 else ''}  ·  Exported {_export_date}",
              new_x="LMARGIN", new_y="NEXT")
    _pdf.set_text_color(*_PT); _pdf.ln(8)

    # Table
    _y_ts = _pdf.get_y()
    _pdf_col_headers(_pdf, _cols, _col_w_pdf, _hdr_h)
    _page_bottom = _pdf.h - 16

    for _ri, (_, _dr) in enumerate(hdf.iterrows()):
        _rv = [str(_dr.get(c, "") or "") for c in _cols]
        _ml = max(len(_pdf.multi_cell(w, _line_h, v, dry_run=True, output="LINES"))
                  for v, w in zip(_rv, _col_w_pdf))
        _rh = max(_ml * _line_h, _line_h)

        if _pdf.get_y() + _rh > _page_bottom:
            _pdf.set_draw_color(*_PBR); _pdf.set_line_width(0.3)
            _pdf.line(_pdf.l_margin, _y_ts, _pdf.l_margin, _pdf.get_y())
            _pdf.line(_pdf.l_margin + _total_w, _y_ts, _pdf.l_margin + _total_w, _pdf.get_y())
            _pdf.add_page()
            _pdf.set_font("DejaVu", size=8.5); _pdf.set_text_color(*_PS)
            _pdf.cell(0, 5, f"{title} — continued", new_x="LMARGIN", new_y="NEXT"); _pdf.ln(3)
            _y_ts = _pdf.get_y()
            _pdf_col_headers(_pdf, _cols, _col_w_pdf, _hdr_h)
            _page_bottom = _pdf.h - 16

        _x0 = _pdf.l_margin; _y0 = _pdf.get_y()
        if _ri % 2 == 1:
            _pdf.set_fill_color(*_PA); _pdf.rect(_x0, _y0, _total_w, _rh, "F")

        _pdf.set_font("DejaVu", size=9); _pdf.set_text_color(*_PT)
        for v, w in zip(_rv, _col_w_pdf):
            _pdf.multi_cell(w, _line_h, f"   {v}", border=0, fill=False,
                            new_x="RIGHT", new_y="TOP", max_line_height=_line_h, align="L")

        _pdf.set_draw_color(*_PBL); _pdf.set_line_width(0.15)
        _pdf.line(_x0, _y0 + _rh, _x0 + _total_w, _y0 + _rh)
        _pdf.set_y(_y0 + _rh)

    _pdf.set_draw_color(*_PBR); _pdf.set_line_width(0.3)
    _ye = _pdf.get_y()
    _pdf.line(_pdf.l_margin, _y_ts, _pdf.l_margin, _ye)
    _pdf.line(_pdf.l_margin + _total_w, _y_ts, _pdf.l_margin + _total_w, _ye)
    _pdf.line(_pdf.l_margin, _ye, _pdf.l_margin + _total_w, _ye)

    pdf_bytes = bytes(_pdf.output())

    # ── Download card grid (3×2) ───────────────────────────────────
    dl1, dl2 = st.columns(2)

    with dl1:
        st.markdown(
            '<div class="dl-card"><div class="dl-title">📄 CSV</div>'
            '<div class="dl-desc">UTF-8 with BOM — opens cleanly in Excel, Google Sheets, or any data tool.</div></div>',
            unsafe_allow_html=True,
        )
        st.download_button("Download CSV", csv_bytes, file_name=f"{file_stem}.csv",
                           mime="text/csv", use_container_width=True, key=f"{key_prefix}_csv")

    with dl2:
        st.markdown(
            '<div class="dl-card"><div class="dl-title">📋 TSV</div>'
            '<div class="dl-desc">Tab-separated — paste into spreadsheets or import into Anki.</div></div>',
            unsafe_allow_html=True,
        )
        st.download_button("Download TSV", tsv_bytes, file_name=f"{file_stem}.tsv",
                           mime="text/tab-separated-values", use_container_width=True, key=f"{key_prefix}_tsv")

    dl3, dl4 = st.columns(2)

    with dl3:
        st.markdown(
            '<div class="dl-card"><div class="dl-title">📊 XLSX</div>'
            '<div class="dl-desc">Excel workbook with all columns — use as a reference or re-import.</div></div>',
            unsafe_allow_html=True,
        )
        st.download_button("Download XLSX", xlsx_bytes, file_name=f"{file_stem}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           use_container_width=True, key=f"{key_prefix}_xlsx")

    with dl4:
        st.markdown(
            '<div class="dl-card"><div class="dl-title">🖹 Markdown</div>'
            '<div class="dl-desc">GFM pipe table — paste into Obsidian, a README, or any Markdown editor.</div></div>',
            unsafe_allow_html=True,
        )
        st.download_button("Download Markdown", md_bytes, file_name=f"{file_stem}.md",
                           mime="text/markdown", use_container_width=True, key=f"{key_prefix}_md")

    dl5, dl6 = st.columns(2)

    with dl5:
        st.markdown(
            '<div class="dl-card"><div class="dl-title">📝 DOCX</div>'
            '<div class="dl-desc">Formatted Word document — landscape A4, styled table, ready to share or print.</div></div>',
            unsafe_allow_html=True,
        )
        st.download_button("Download DOCX", docx_bytes, file_name=f"{file_stem}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True, key=f"{key_prefix}_docx")

    with dl6:
        st.markdown(
            '<div class="dl-card"><div class="dl-title">🖨 PDF</div>'
            '<div class="dl-desc">Formatted PDF — landscape A4, styled table, ready to print or share.</div></div>',
            unsafe_allow_html=True,
        )
        st.download_button("Download PDF", pdf_bytes, file_name=f"{file_stem}.pdf",
                           mime="application/pdf", use_container_width=True, key=f"{key_prefix}_pdf")


# ---------------------------------------------------------------------------
# Guard
# ---------------------------------------------------------------------------

db_path = get_db_path()
if not Path(db_path).exists():
    st.title("📋 Report")
    st.error("Database not found. Open ⚙️ Database settings in the sidebar and click **Initialize DB**.")
    st.stop()

title_slot = st.empty()   # filled with h2 heading once filter_label is known
pill_slot  = st.empty()   # filled with filter pill

# ---------------------------------------------------------------------------
# Controls
# ---------------------------------------------------------------------------

with st.container(border=True):
    r1c1, r1c2, r1c3, r1c4 = st.columns([3, 2, 2, 3])

    days     = r1c1.slider("Look-back window (days)",        min_value=7,  max_value=90, value=30, step=7)
    worst_n  = r1c2.slider("Worst items",                    min_value=5,  max_value=50, value=20, step=5)
    missed_n = r1c3.slider("Most-missed (all-time)",         min_value=5,  max_value=50, value=20, step=5)

    con = open_db(db_path)
    try:
        sources = list_sources(con)
    finally:
        con.close()

    source_options = ["Auto (latest pipeline)", "All sources"] + sources
    source_choice  = r1c4.selectbox("Vocabulary source", source_options, label_visibility="visible")

cutoff = cutoff_iso(days)

# Resolve source
if source_choice == "Auto (latest pipeline)":
    con = open_db(db_path)
    try:
        resolved_source = storage.get_latest_pipeline_source(con)
    finally:
        con.close()
    source, source_prefix, pipeline_only = resolved_source, None, True
elif source_choice == "All sources":
    source, source_prefix, pipeline_only = None, None, False
else:
    source, source_prefix, pipeline_only = source_choice, None, False

# ---------------------------------------------------------------------------
# Summary stats
# ---------------------------------------------------------------------------

con = open_db(db_path)
try:
    stats = storage.query_stats(
        con, cutoff,
        source=source, source_prefix=source_prefix,
        default_pipeline_only=pipeline_only,
    )
finally:
    con.close()

filter_label = stats.get("filter_label", source_choice)
_c = get_plotly_colors()   # theme-aware colour palette for all charts this render

# ── Page header — fills the placeholders reserved at the top of the page ──────
title_slot.markdown("## 📋 Report")
pill_slot.markdown(
    f'<div style="margin-top:-0.75rem;margin-bottom:0.75rem;">'
    f'<span class="filter-pill">Last {days} days · {filter_label}</span>'
    f'</div>',
    unsafe_allow_html=True,
)

# ── Metric cards ──────────────────────────────────────────────────────────────
m1, m2, m3, m4 = st.columns(4)

def _card(col, num, lbl, sub=""):
    col.markdown(
        f'<div class="stat-card">'
        f'<div class="sc-num">{num}</div>'
        f'<div class="sc-lbl">{lbl}</div>'
        f'<div class="sc-sub">{sub if sub else "&nbsp;"}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )

_card(m1, stats["vocab_count"],          "Vocabulary",   "words in source")
_card(m2, stats["attempts_count"],       "Attempts",     f"last {days} days")
_card(m3, fmt_rate(stats["accuracy"]),   "Accuracy",     "window average")
_card(m4, fmt_ts(stats["last_seen"]),    "Last Practice","")

# ---------------------------------------------------------------------------
# Section 1 — Worst items (window)
# ---------------------------------------------------------------------------

_worst_hdr_left, _worst_hdr_right = st.columns([4, 1])
_worst_hdr_left.markdown(
    f'<div class="section-lbl">Worst {worst_n} Items — Last {days} Days</div>',
    unsafe_allow_html=True,
)
include_unattempted = _worst_hdr_right.checkbox(
    "Include unattempted",
    value=False,
    help="When checked, words you haven't practised yet are included (shown as 0 %).",
)

con = open_db(db_path)
try:
    worst_rows = storage.query_worst_items(
        con, cutoff, worst_n,
        source=source, source_prefix=source_prefix,
        default_pipeline_only=pipeline_only,
        min_attempts=0 if include_unattempted else 1,
    )
finally:
    con.close()

if worst_rows:
    worst_df = pd.DataFrame(worst_rows)

    # ── Horizontal bar chart ──────────────────────────────────────────────────
    chart_df = worst_df.sort_values("acc_window", ascending=True).head(15)
    labels   = chart_df["de_display"].fillna(chart_df["en"]).str[:28]
    accs     = (chart_df["acc_window"] * 100).round(0)
    colours  = [
        "#FC8181" if a < 40 else ("#ECC94B" if a < 70 else "#48BB78")
        for a in accs
    ]

    fig = go.Figure(go.Bar(
        x=accs,
        y=labels,
        orientation="h",
        marker=dict(color=colours, line=dict(width=0)),
        text=[f"{a:.0f}%" for a in accs],
        textposition="outside",
        textfont=dict(size=11),
        hovertemplate=(
            "<b>%{y}</b><br>"
            "Accuracy: %{x:.0f}%<br>"
            "Attempts: %{customdata}<extra></extra>"
        ),
        customdata=chart_df["attempts_window"],
    ))
    fig.update_layout(**build_plotly_layout(_c,
        xaxis=dict(range=[0, 118], ticksuffix="%", zeroline=False),
        yaxis=dict(autorange="reversed", showgrid=False, zeroline=False),
        bargap=0.28,
        height=max(220, len(chart_df) * 28),
        margin=dict(r=52, t=8, b=4),
    ))
    st.plotly_chart(fig, width="stretch", theme=None)

    # ── Table ─────────────────────────────────────────────────────────────────
    tbl = worst_df.rename(columns={
        "de_display":       "German",
        "en":               "English",
        "acc_window":       "Accuracy",
        "attempts_window":  "Attempts",
        "near_miss_window": "Near-miss",
        "last_seen":        "Last seen",
    }).copy()
    tbl["Last seen"] = tbl["Last seen"].apply(lambda x: x[:10] if x else "—")

    st.dataframe(
        tbl[["German", "English", "Accuracy", "Attempts", "Near-miss", "Last seen"]],
        width="stretch",
        hide_index=True,
        column_config={
            "Accuracy": st.column_config.ProgressColumn(
                "Accuracy",
                min_value=0,
                max_value=1,
                format="%.0f%%",
                help="Window accuracy rate",
            ),
            "German":    st.column_config.TextColumn("German",    width="medium"),
            "English":   st.column_config.TextColumn("English",   width="medium"),
            "Last seen": st.column_config.TextColumn("Last seen", width="small"),
        },
    )

    # ── Export ────────────────────────────────────────────────────────────────
    _build_exports(
        tbl[["German", "English", "Accuracy", "Attempts", "Near-miss", "Last seen"]],
        title=f"Worst {worst_n} Items — Last {days} Days",
        file_stem=f"worst_items_{days}d",
        pct_cols=["Accuracy"],
        key_prefix="worst",
    )
else:
    if include_unattempted:
        st.info("No vocabulary items found for the selected source.")
    else:
        st.info("No practised items in the selected window. Check **Include unattempted** to see all words.")

# ---------------------------------------------------------------------------
# Section 2 — Most missed (all-time)
# ---------------------------------------------------------------------------

st.markdown(
    f'<div class="section-lbl">Most Missed — All-Time Top {missed_n}</div>',
    unsafe_allow_html=True,
)

con = open_db(db_path)
try:
    missed_rows = storage.query_most_missed_alltime(
        con, top_n=missed_n,
        source=source, source_prefix=source_prefix,
        default_pipeline_only=pipeline_only,
    )
finally:
    con.close()

if missed_rows:
    missed_df = pd.DataFrame(missed_rows).rename(columns={
        "de_display":     "German",
        "en":             "English",
        "miss_count":     "Times missed",
        "total_attempts": "Total attempts",
        "acc_alltime":    "All-time accuracy",
    })

    # ── Scatter: missed count vs all-time accuracy ────────────────────────────
    scatter_df = missed_df.copy()
    scatter_df["label"] = scatter_df["German"].fillna(scatter_df["English"]).str[:24]
    scatter_df["color"] = scatter_df["All-time accuracy"].apply(
        lambda a: "#FC8181" if a < 0.4 else ("#ECC94B" if a < 0.7 else "#48BB78")
    )

    fig_scatter = go.Figure(go.Scatter(
        x=scatter_df["All-time accuracy"] * 100,
        y=scatter_df["Times missed"],
        mode="markers+text",
        text=scatter_df["label"],
        textposition="top center",
        textfont=dict(size=9),
        marker=dict(
            color=scatter_df["color"].tolist(),
            size=scatter_df["Total attempts"].apply(lambda n: max(8, min(n * 1.5, 24))).tolist(),
            opacity=0.8,
            line=dict(color="rgba(0,0,0,0.3)", width=1),
        ),
        hovertemplate=(
            "<b>%{text}</b><br>"
            "All-time accuracy: %{x:.0f}%<br>"
            "Times missed: %{y}<extra></extra>"
        ),
        customdata=scatter_df["Total attempts"],
        showlegend=False,
    ))
    fig_scatter.add_vline(
        x=70,
        line=dict(color="rgba(72,187,120,0.25)", dash="dot", width=1.5),
        annotation_text="70% target",
        annotation_position="top right",
        annotation_font=dict(size=10, color="rgba(72,187,120,0.5)"),
    )
    fig_scatter.update_layout(**build_plotly_layout(_c,
        xaxis=dict(title="All-time accuracy", ticksuffix="%", range=[-2, 105], zeroline=False),
        yaxis=dict(title="Times missed", zeroline=False),
        height=320,
        margin=dict(l=4, r=4, t=8, b=4),
    ))
    st.plotly_chart(fig_scatter, width="stretch", theme=None)
    st.markdown(
        '<div style="font-size:0.7rem;color:color-mix(in srgb, var(--text-color) 22%, transparent);margin-top:-0.5rem;margin-bottom:0.75rem;">'
        'Bubble size ∝ total attempts. Words in the top-left corner are both frequently missed '
        'and have low accuracy — highest priority for review.</div>',
        unsafe_allow_html=True,
    )

    # ── Table ─────────────────────────────────────────────────────────────────
    st.dataframe(
        missed_df[["German", "English", "Times missed", "Total attempts", "All-time accuracy"]],
        width="stretch",
        hide_index=True,
        column_config={
            "All-time accuracy": st.column_config.ProgressColumn(
                "All-time accuracy",
                min_value=0,
                max_value=1,
                format="%.0f%%",
            ),
            "German":         st.column_config.TextColumn("German",         width="medium"),
            "English":        st.column_config.TextColumn("English",        width="medium"),
            "Times missed":   st.column_config.NumberColumn("Times missed", width="small"),
            "Total attempts": st.column_config.NumberColumn("Attempts",     width="small"),
        },
    )

    # ── Export ────────────────────────────────────────────────────────────────
    _build_exports(
        missed_df[["German", "English", "Times missed", "Total attempts", "All-time accuracy"]],
        title=f"Most Missed — All-Time Top {missed_n}",
        file_stem="most_missed_alltime",
        pct_cols=["All-time accuracy"],
        key_prefix="missed",
    )
else:
    st.info("No items with enough practice data yet (minimum 3 attempts required).")

"""
pages/1_Vocab_Export.py — Paste German words → GPT-4o enrichment → download.
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from datetime import datetime, timezone
import pandas as pd
import streamlit as st

# Ensure project root is on sys.path so src/ is importable
_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from src.vocab_export_core import clean_text, get_vocabulary_data
from ui_utils import get_db_path, open_db
from german_pipeline import ingest_export


# ── Shared DB import helper ───────────────────────────────────────────────────
def _do_db_import(df: "pd.DataFrame", source_label: str) -> None:
    """Upsert *df* rows into the practice DB under pipeline:<source_label>."""
    if not source_label.strip():
        st.error("Please enter a source label.")
        return
    rows = [
        {
            "de":             row.get("Deutsch", ""),
            "de_mit_artikel": row.get("Deutsch mit Artikel", ""),
            "en":             row.get("Englisch", ""),
            "af":             row.get("Afrikaans", ""),
            "notes":          row.get("Wortart / Genus / Hinweise", ""),
        }
        for _, row in df.iterrows()
    ]
    db_path = get_db_path()
    con = open_db(db_path)
    try:
        full_source = f"pipeline:{source_label.strip()}"
        inserted, updated = ingest_export.upsert_vocab_items(con, rows, full_source)
        con.commit()
    except Exception as e:
        con.close()
        st.error(f"Import failed: {e}")
        return
    finally:
        con.close()

    st.success(f"**{full_source}** saved to practice DB.")
    st.markdown(f"""
<div class="result-row">
  <div class="result-pill pill-inserted">
    <div class="rp-num">{inserted}</div>
    <div class="rp-lbl">Inserted</div>
  </div>
  <div class="result-pill pill-updated">
    <div class="rp-num">{updated}</div>
    <div class="rp-lbl">Updated</div>
  </div>
</div>
""", unsafe_allow_html=True)


# ── Page-scoped CSS ──────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Pipeline field chips ───────────────────────────────────────── */
.field-row {
    display: flex;
    flex-wrap: wrap;
    gap: 0.45rem;
    margin: 0.5rem 0 0.25rem;
}
.field-chip {
    display: inline-flex;
    align-items: center;
    gap: 0.3rem;
    background: color-mix(in srgb, var(--text-color) 5%, transparent);
    border: 1px solid color-mix(in srgb, var(--text-color) 10%, transparent);
    border-radius: 20px;
    padding: 0.2rem 0.65rem;
    font-size: 0.75rem;
    color: color-mix(in srgb, var(--text-color) 65%, transparent);
}
.field-chip .dot {
    width: 6px; height: 6px;
    border-radius: 50%;
    flex-shrink: 0;
}
/* ── Anki card mockup ───────────────────────────────────────────── */
.anki-card {
    border-radius: 10px;
    overflow: hidden;
    border: 1px solid color-mix(in srgb, var(--text-color) 10%, transparent);
    max-width: 420px;
    margin: 0.5rem auto;
}
.anki-front {
    background: linear-gradient(135deg, #1a365d 0%, #2c5282 100%);
    padding: 1.1rem 1.4rem 0.9rem;
    text-align: center;
}
.anki-front .af-label {
    font-size: 0.6rem;
    font-weight: 700;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: rgba(255,255,255,0.35);
    margin-bottom: 0.4rem;
}
.anki-front .af-word {
    font-size: 1.35rem;
    font-weight: 700;
    color: #fff;
}
.anki-front .af-hint {
    font-size: 0.75rem;
    color: rgba(255,255,255,0.45);
    margin-top: 0.25rem;
}
.anki-back {
    background: color-mix(in srgb, var(--text-color) 4%, transparent);
    padding: 0.8rem 1.4rem;
    text-align: center;
    border-top: 1px solid color-mix(in srgb, var(--text-color) 8%, transparent);
}
.anki-back .ab-label {
    font-size: 0.6rem;
    font-weight: 700;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: color-mix(in srgb, var(--text-color) 30%, transparent);
    margin-bottom: 0.3rem;
}
.anki-back .ab-word {
    font-size: 1.1rem;
    font-weight: 600;
    color: color-mix(in srgb, var(--text-color) 88%, transparent);
}
/* ── Download cards ─────────────────────────────────────────────── */
.dl-card {
    background: color-mix(in srgb, var(--text-color) 4%, transparent);
    border: 1px solid color-mix(in srgb, var(--text-color) 10%, transparent);
    border-radius: 8px;
    padding: 0.9rem 1rem 0.6rem;
    margin-bottom: 0.5rem;
}
.dl-card .dl-title {
    font-weight: 600;
    font-size: 0.9rem;
    color: color-mix(in srgb, var(--text-color) 85%, transparent);
    margin-bottom: 0.15rem;
}
.dl-card .dl-desc {
    font-size: 0.73rem;
    color: color-mix(in srgb, var(--text-color) 38%, transparent);
    margin-bottom: 0.65rem;
    line-height: 1.4;
}
/* ── Result metrics row ─────────────────────────────────────────── */
.result-row {
    display: flex;
    gap: 0.75rem;
    margin-top: 0.75rem;
}
.result-pill {
    flex: 1;
    border-radius: 8px;
    padding: 0.7rem 0.6rem 0.5rem;
    text-align: center;
    border: 1px solid color-mix(in srgb, var(--text-color) 6%, transparent);
}
.result-pill .rp-num {
    font-size: 1.6rem;
    font-weight: 800;
    line-height: 1;
}
.result-pill .rp-lbl {
    font-size: 0.62rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    margin-top: 0.25rem;
    opacity: 0.55;
}
.pill-inserted { background: rgba(72,187,120,0.12);  color: #68D391; }
.pill-updated  { background: rgba(246,173,85,0.12);  color: #F6AD55; }
/* ── Cached badge ───────────────────────────────────────────────── */
.cache-badge {
    display: inline-block;
    background: rgba(246,173,85,0.15);
    border: 1px solid rgba(246,173,85,0.3);
    color: #F6AD55;
    font-size: 0.7rem;
    font-weight: 600;
    letter-spacing: 0.06em;
    padding: 0.15rem 0.55rem;
    border-radius: 20px;
    vertical-align: middle;
    margin-left: 0.5rem;
}
</style>
""", unsafe_allow_html=True)

# ── Page header ──────────────────────────────────────────────────────────────
st.markdown("## 📝 Vocab Export")
st.markdown(
    '<p style="color:color-mix(in srgb, var(--text-color) 45%, transparent);margin-top:-0.4rem;font-size:0.9rem;">'
    "Paste raw German words or phrases — GPT-4o enriches each one and generates "
    "Anki-ready flashcard exports."
    "</p>",
    unsafe_allow_html=True,
)

# ── What GPT-4o generates ────────────────────────────────────────────────────
st.markdown(
    '<div class="field-row">'
    '<div class="field-chip"><span class="dot" style="background:#63B3ED"></span>Deutsch</div>'
    '<div class="field-chip"><span class="dot" style="background:#68D391"></span>Deutsch mit Artikel</div>'
    '<div class="field-chip"><span class="dot" style="background:#F6AD55"></span>Englisch</div>'
    '<div class="field-chip"><span class="dot" style="background:#FC8181"></span>Afrikaans</div>'
    '<div class="field-chip"><span class="dot" style="background:#B794F4"></span>Wortart / Genus / Hinweise</div>'
    '</div>',
    unsafe_allow_html=True,
)

st.markdown('<div style="height:0.75rem"></div>', unsafe_allow_html=True)

# ── Input area ───────────────────────────────────────────────────────────────
raw_text = st.text_area(
    "Paste German words or phrases — one per line",
    height=220,
    placeholder="die Autobahn\nschreiben\nins kalte Wasser springen\nder Kürbis\n...",
    label_visibility="visible",
)

# Word count hint
if raw_text.strip():
    phrase_count = len([l for l in raw_text.strip().splitlines() if l.strip()])
    st.markdown(
        f'<span style="font-size:0.78rem; color:color-mix(in srgb, var(--text-color) 35%, transparent);">'
        f'{phrase_count} phrase{"s" if phrase_count != 1 else ""} detected</span>',
        unsafe_allow_html=True,
    )

st.markdown('<div style="height:0.4rem"></div>', unsafe_allow_html=True)

# ── Controls row ─────────────────────────────────────────────────────────────
ctrl_left, ctrl_right = st.columns([3, 2])

process_clicked = ctrl_left.button(
    "✨  Enrich with GPT-4o",
    disabled=not raw_text.strip(),
    type="primary",
    use_container_width=True,
)

auto_import = ctrl_right.toggle(
    "Auto-import to DB after processing",
    value=False,
    help="When on, results are sent straight to the practice database once enrichment finishes.",
)

if auto_import:
    source_label_auto = st.text_input(
        "Source label",
        value="vocab_export",
        help="Stored as pipeline:<label> in the DB.",
        placeholder="e.g. session_01",
    )

# ── Processing ───────────────────────────────────────────────────────────────
if process_clicked and raw_text.strip():
    input_hash = hash(raw_text.strip())

    if st.session_state.get("export_input_hash") == input_hash:
        from_cache = True
    else:
        from_cache = False
        phrases = clean_text(raw_text)

        with st.spinner(f"Calling GPT-4o for {len(phrases)} phrase(s)…"):
            try:
                vocab_list = get_vocabulary_data(phrases)
            except Exception as e:
                st.error(f"API error: {e}")
                st.stop()

        df = pd.DataFrame(vocab_list)
        df.rename(columns={
            "deutsch":             "Deutsch",
            "deutsch_mit_artikel": "Deutsch mit Artikel",
            "englisch":            "Englisch",
            "afrikaans":           "Afrikaans",
            "hinweise":            "Wortart / Genus / Hinweise",
        }, inplace=True)

        st.session_state["export_df"]         = df
        st.session_state["export_input_hash"] = input_hash
        st.session_state["export_from_cache"] = False

    if from_cache:
        st.session_state["export_from_cache"] = True

    # Auto-import runs immediately after enrichment (not on slider re-runs)
    if auto_import and not from_cache:
        _label = source_label_auto if "source_label_auto" in dir() else "vocab_export"
        _do_db_import(st.session_state["export_df"], _label)

# ── Results (rendered whenever data exists in session state) ──────────────────
if "export_df" in st.session_state:
    df         = st.session_state["export_df"]
    from_cache = st.session_state.get("export_from_cache", False)

    # ── Section heading ──────────────────────────────────────────────────
    cached_badge = (
        '<span class="cache-badge">⚡ CACHED</span>' if from_cache else ""
    )
    st.markdown(
        f'<div style="height:1rem"></div>'
        f'<h3 style="margin-bottom:0.2rem">Enriched Vocabulary {cached_badge}</h3>'
        f'<p style="color:color-mix(in srgb, var(--text-color) 40%, transparent); font-size:0.83rem; margin-top:0;">'
        f'{len(df)} item{"s" if len(df) != 1 else ""} enriched</p>',
        unsafe_allow_html=True,
    )

    # ── Result tabs ──────────────────────────────────────────────────────
    tab_vocab, tab_anki, tab_export = st.tabs([
        "📋  Vocabulary table",
        "🃏  Anki card preview",
        "⬇️  Download & import",
    ])

    # ── Tab 1: Vocabulary table ──────────────────────────────────────────
    with tab_vocab:
        st.dataframe(
            df,
            width="stretch",
            hide_index=True,
            column_config={
                "Deutsch":                 st.column_config.TextColumn("Deutsch",       width="medium"),
                "Deutsch mit Artikel":     st.column_config.TextColumn("Mit Artikel",   width="medium"),
                "Englisch":                st.column_config.TextColumn("Englisch",       width="medium"),
                "Afrikaans":               st.column_config.TextColumn("Afrikaans",      width="medium"),
                "Wortart / Genus / Hinweise": st.column_config.TextColumn("Hinweise",   width="large"),
            },
        )

    # ── Tab 2: Anki card preview ─────────────────────────────────────────
    with tab_anki:
        st.markdown(
            '<p style="color:color-mix(in srgb, var(--text-color) 40%, transparent); font-size:0.82rem; margin-bottom:0.75rem;">'
            "Each card has the <strong>English meaning + grammar notes</strong> on the front "
            "and the <strong>German word with article</strong> on the back.</p>",
            unsafe_allow_html=True,
        )

        preview_idx = st.slider(
            "Preview card",
            min_value=1,
            max_value=len(df),
            value=1,
            key="anki_preview_slider",
            label_visibility="visible",
        ) - 1

        row = df.iloc[preview_idx]
        front_hint  = row.get("Wortart / Genus / Hinweise", "")
        front_en    = row.get("Englisch", "")
        back_text   = row.get("Deutsch mit Artikel", row.get("Deutsch", ""))

        st.markdown(f"""
<div class="anki-card">
  <div class="anki-front">
    <div class="af-label">Front</div>
    <div class="af-word">{front_en}</div>
    <div class="af-hint">{front_hint}</div>
  </div>
  <div class="anki-back">
    <div class="ab-label">Back</div>
    <div class="ab-word">{back_text}</div>
  </div>
</div>
""", unsafe_allow_html=True)

    # ── Tab 3: Download & import ─────────────────────────────────────────
    with tab_export:
        # ── Build export payloads ──────────────────────────────────────
        # XLSX — full vocabulary
        xlsx_buf = io.BytesIO()
        df.to_excel(xlsx_buf, index=False, engine="openpyxl")
        xlsx_bytes = xlsx_buf.getvalue()

        # CSV — full vocabulary (UTF-8 with BOM for Excel compat)
        csv_buf = io.StringIO()
        df.to_csv(csv_buf, index=False, encoding="utf-8")
        csv_bytes = ("\ufeff" + csv_buf.getvalue()).encode("utf-8")

        # TSV — Anki import (two-column, headerless)
        anki_df = df[["Deutsch mit Artikel", "Englisch", "Wortart / Genus / Hinweise"]].copy()
        anki_df["Front"] = anki_df["Englisch"].fillna("") + " — " + anki_df["Wortart / Genus / Hinweise"].fillna("")
        anki_df["Back"]  = anki_df["Deutsch mit Artikel"]
        anki_export = anki_df[["Front", "Back"]]
        tsv_buf = io.StringIO()
        anki_export.to_csv(tsv_buf, sep="\t", index=False, header=False, encoding="utf-8")
        tsv_bytes = tsv_buf.getvalue().encode("utf-8")

        # DOCX — polished landscape vocabulary document
        import docx as _docx
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _Ox

        # ── DOCX design tokens ───────────────────────────────────────────────
        # Muted slate-blue palette — professional, restrained
        _DH_BG   = "3B5998"                        # header fill
        _DH_TEXT = RGBColor(0xFF, 0xFF, 0xFF)       # header text
        _DT_CLR  = RGBColor(0x2D, 0x3A, 0x4A)      # title text
        _DS_CLR  = RGBColor(0x8B, 0x95, 0xA5)      # subtitle
        _DB_CLR  = RGBColor(0x2D, 0x3A, 0x4A)      # body text

        # ── DOCX styling helpers ─────────────────────────────────────────────
        def _d_shade(cell, hx):
            shd = _Ox("w:shd")
            shd.set(_qn("w:fill"), hx)
            shd.set(_qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shd)

        def _d_tbl_borders(tbl):
            """Soft outer border + very faint row lines. No inner verticals."""
            tblPr = tbl._tbl.tblPr
            def _b(tag, color, sz=4, val="single"):
                b = _Ox(f"w:{tag}")
                b.set(_qn("w:val"), val)
                b.set(_qn("w:sz"), str(sz))
                b.set(_qn("w:color"), color)
                return b
            bdr = _Ox("w:tblBorders")
            bdr.append(_b("top",     "CFD8DC", sz=4))
            bdr.append(_b("left",    "CFD8DC", sz=2))
            bdr.append(_b("bottom",  "CFD8DC", sz=4))
            bdr.append(_b("right",   "CFD8DC", sz=2))
            bdr.append(_b("insideH", "E8ECF0", sz=2))
            bdr.append(_b("insideV", "none",   sz=0, val="none"))
            tblPr.append(bdr)

        def _d_cell_margins(tbl, top=100, start=140, bottom=100, end=140):
            """Generous table-level cell padding (in dxa; 20 dxa = 1 pt)."""
            tblPr = tbl._tbl.tblPr
            mar = _Ox("w:tblCellMar")
            for side, val in [("top", top), ("start", start),
                              ("bottom", bottom), ("end", end)]:
                n = _Ox(f"w:{side}")
                n.set(_qn("w:w"), str(val))
                n.set(_qn("w:type"), "dxa")
                mar.append(n)
            tblPr.append(mar)

        def _d_para_rule(para, color="B0BEC5"):
            """Paragraph bottom border — decorative rule."""
            pPr = para._p.get_or_add_pPr()
            pBdr = _Ox("w:pBdr")
            b = _Ox("w:bottom")
            b.set(_qn("w:val"), "single")
            b.set(_qn("w:sz"), "6")
            b.set(_qn("w:color"), color)
            b.set(_qn("w:space"), "8")
            pBdr.append(b)
            pPr.append(pBdr)

        def _d_repeat_header(tbl):
            """Mark first row as repeating header on page breaks."""
            trPr = tbl.rows[0]._tr.get_or_add_trPr()
            trPr.append(_Ox("w:tblHeader"))

        # ── Document ─────────────────────────────────────────────────────────
        _export_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
        _doc = _docx.Document()
        _doc.core_properties.title   = "German Vocabulary Export"
        _doc.core_properties.subject = "Language learning flashcard vocabulary"

        # Page setup — A4 landscape with generous margins
        _section = _doc.sections[0]
        _section.page_width    = Cm(29.7)
        _section.page_height   = Cm(21.0)
        _section.left_margin   = Cm(2.5)
        _section.right_margin  = Cm(2.5)
        _section.top_margin    = Cm(2.0)
        _section.bottom_margin = Cm(1.8)

        # Title — large, dark slate
        _tp = _doc.add_paragraph()
        _tp.paragraph_format.space_before = Pt(0)
        _tp.paragraph_format.space_after  = Pt(4)
        _tr = _tp.add_run("German Vocabulary Export")
        _tr.bold      = True
        _tr.font.size = Pt(20)
        _tr.font.color.rgb = _DT_CLR
        _d_para_rule(_tp)

        # Subtitle — item count (left) + export date (right-aligned tab stop)
        _sp = _doc.add_paragraph()
        _sp.paragraph_format.space_before = Pt(8)
        _sp.paragraph_format.space_after  = Pt(20)
        # Right-aligned tab stop at the right margin (usable width ≈ 24.7 cm)
        _usable_w = _section.page_width - _section.left_margin - _section.right_margin
        _tab_stops = _sp.paragraph_format.tab_stops
        _tab_stops.add_tab_stop(_usable_w, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        _sr_left = _sp.add_run(
            f"{len(df)} item{'s' if len(df) != 1 else ''} enriched by GPT-4o"
        )
        _sr_left.font.size = Pt(9.5)
        _sr_left.font.color.rgb = _DS_CLR
        _sp.add_run("\t")  # tab character → jumps to right-aligned stop
        _sr_right = _sp.add_run(f"Exported {_export_date}")
        _sr_right.font.size = Pt(9.5)
        _sr_right.font.color.rgb = _DS_CLR

        # Table — widths tuned for landscape A4 (usable ≈ 24.7 cm)
        _cols = list(df.columns)
        _D_COL_W = {
            "Deutsch":                    Cm(4.6),
            "Deutsch mit Artikel":        Cm(5.0),
            "Englisch":                   Cm(5.0),
            "Afrikaans":                  Cm(4.6),
            "Wortart / Genus / Hinweise": Cm(5.5),
        }
        _tbl = _doc.add_table(rows=1, cols=len(_cols))
        _tbl.style     = "Table Grid"
        _tbl.autofit   = False
        _tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _d_cell_margins(_tbl)
        _d_tbl_borders(_tbl)
        for i, col_name in enumerate(_cols):
            _tbl.columns[i].width = _D_COL_W.get(col_name, Cm(4.0))

        # Header row — muted-blue fill, white bold text, extra height
        _hdr_cells = _tbl.rows[0].cells
        for i, col_name in enumerate(_cols):
            _hdr_cells[i].text = col_name
            _hdr_cells[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            _d_shade(_hdr_cells[i], _DH_BG)
            _run = _hdr_cells[i].paragraphs[0].runs[0]
            _run.bold      = True
            _run.font.size = Pt(9)
            _run.font.color.rgb = _DH_TEXT
            _hdr_cells[i].paragraphs[0].paragraph_format.space_before = Pt(3)
            _hdr_cells[i].paragraphs[0].paragraph_format.space_after  = Pt(3)
        _d_repeat_header(_tbl)

        # Data rows — alternating, comfortable spacing
        for _row_idx, (_, _data_row) in enumerate(df.iterrows()):
            _row_cells = _tbl.add_row().cells
            for i, col_name in enumerate(_cols):
                _row_cells[i].text = str(_data_row.get(col_name, "") or "")
                if _row_idx % 2 == 1:
                    _d_shade(_row_cells[i], "F5F7FA")
                _run = _row_cells[i].paragraphs[0].runs[0]
                _run.font.size = Pt(9.5)
                _run.font.color.rgb = _DB_CLR
                _row_cells[i].paragraphs[0].paragraph_format.space_before = Pt(1)
                _row_cells[i].paragraphs[0].paragraph_format.space_after  = Pt(1)

        docx_buf = io.BytesIO()
        _doc.save(docx_buf)
        docx_bytes = docx_buf.getvalue()

        # Markdown — GFM pipe table
        def _md_cell(val: str) -> str:
            """Escape pipe characters so the GFM table stays intact."""
            return val.replace("|", "\\|").replace("\n", " ")

        _cols = list(df.columns)
        _md_parts = [
            "# German Vocabulary Export\n",
            f"_{len(df)} item{'s' if len(df) != 1 else ''} enriched by GPT-4o_\n",
            "",
            "| " + " | ".join(_md_cell(c) for c in _cols) + " |",
            "| " + " | ".join("---" for _ in _cols) + " |",
        ]
        for _, data_row in df.iterrows():
            _md_parts.append(
                "| " + " | ".join(
                    _md_cell(str(data_row.get(c, "") or "")) for c in _cols
                ) + " |"
            )
        md_bytes = "\n".join(_md_parts).encode("utf-8")

        # PDF — polished landscape table via fpdf2
        # Uses bundled DejaVu Sans (assets/fonts/) — cross-platform, open license,
        # full Unicode coverage for German, Afrikaans, and arbitrary GPT-4o output.
        from fpdf import FPDF

        _FONTS_DIR    = Path(__file__).resolve().parent.parent / "assets" / "fonts"
        _FONT_REGULAR = _FONTS_DIR / "DejaVuSans.ttf"
        _FONT_BOLD    = _FONTS_DIR / "DejaVuSans-Bold.ttf"
        if not _FONT_REGULAR.exists() or not _FONT_BOLD.exists():
            st.error(
                "PDF export requires bundled fonts in `assets/fonts/`. "
                "Expected: `DejaVuSans.ttf` and `DejaVuSans-Bold.ttf`. "
                "See `assets/fonts/` in the repo root."
            )
            st.stop()

        # Design tokens (PDF) — muted slate-blue palette
        _PH   = (59,  89,  152)    # #3B5998 — header fill (muted blue)
        _PHD  = (42,  63,  110)    # darker shade for header accent line
        _PW   = (255, 255, 255)
        _PT   = (45,  58,  74)     # #2D3A4A — title / body text
        _PS   = (139, 149, 165)    # #8B95A5 — subtitle
        _PA   = (245, 247, 250)    # #F5F7FA — alt row bg
        _PBR  = (207, 216, 220)    # #CFD8DC — outer border / footer rule
        _PBL  = (232, 236, 240)    # #E8ECF0 — faint inner row lines

        _PDF_COLS  = list(df.columns)
        # Column widths for landscape A4 — 20 mm margins → usable 257 mm
        _PDF_COL_W = {
            "Deutsch":                    42,
            "Deutsch mit Artikel":        44,
            "Englisch":                   44,
            "Afrikaans":                  40,
            "Wortart / Genus / Hinweise": 87,   # total = 257
        }
        _col_widths = [_PDF_COL_W.get(c, 51) for c in _PDF_COLS]
        _total_w    = sum(_col_widths)
        _line_h     = 7     # mm per wrapped line — spacious
        _hdr_h      = 10    # mm for header row — taller

        def _pdf_col_headers(pdf, cols, widths, h):
            """Muted-blue header band with white text."""
            pdf.set_font("DejaVu", style="B", size=9)
            pdf.set_fill_color(*_PH)
            pdf.set_text_color(*_PW)
            x0, y0 = pdf.l_margin, pdf.get_y()
            pdf.rect(x0, y0, sum(widths), h, "F")
            for col_name, w in zip(cols, widths):
                pdf.multi_cell(
                    w, h, f"  {col_name}",
                    border=0, fill=False,
                    new_x="RIGHT", new_y="TOP",
                    max_line_height=h,
                    align="L",
                )
            pdf.set_y(y0 + h)
            # Accent line — slightly darker than header fill
            pdf.set_draw_color(*_PHD)
            pdf.set_line_width(0.6)
            pdf.line(x0, pdf.get_y(), x0 + sum(widths), pdf.get_y())
            pdf.set_draw_color(*_PBL)
            pdf.set_line_width(0.15)
            pdf.set_text_color(*_PT)

        class _VocabPDF(FPDF):
            def footer(self):
                self.set_y(-14)
                self.set_draw_color(*_PBR)
                self.set_line_width(0.3)
                self.line(self.l_margin, self.get_y(),
                          self.w - self.r_margin, self.get_y())
                self.set_y(-12)
                self.set_font("DejaVu", size=7.5)
                self.set_text_color(*_PS)
                self.cell(0, 5, f"Page {self.page_no()} / {{nb}}", align="R")
                self.set_text_color(*_PT)

        _pdf = _VocabPDF(orientation="L", unit="mm", format="A4")
        _pdf.add_font("DejaVu", style="",  fname=str(_FONT_REGULAR))
        _pdf.add_font("DejaVu", style="B", fname=str(_FONT_BOLD))
        _pdf.alias_nb_pages()
        _pdf.set_margins(left=20, top=18, right=20)
        _pdf.set_auto_page_break(auto=False)
        _pdf.add_page()

        # ── Title block ──────────────────────────────────────────────────────
        _pdf.set_font("DejaVu", style="B", size=20)
        _pdf.set_text_color(*_PT)
        _pdf.cell(0, 12, "German Vocabulary Export",
                  new_x="LMARGIN", new_y="NEXT")

        # Accent rule — full content width
        _pdf.set_draw_color(*_PBR)
        _pdf.set_line_width(0.5)
        _pdf.line(_pdf.l_margin, _pdf.get_y() + 1,
                  _pdf.w - _pdf.r_margin, _pdf.get_y() + 1)
        _pdf.set_line_width(0.15)
        _pdf.ln(5)

        # Subtitle
        _export_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
        _pdf.set_font("DejaVu", size=10)
        _pdf.set_text_color(*_PS)
        _pdf.cell(
            0, 6,
            f"{len(df)} item{'s' if len(df) != 1 else ''} enriched by "
            f"GPT-4o  \u00b7  Exported {_export_date}",
            new_x="LMARGIN", new_y="NEXT",
        )
        _pdf.set_text_color(*_PT)
        _pdf.ln(8)

        # ── Table ────────────────────────────────────────────────────────────
        _y_table_start = _pdf.get_y()
        _pdf_col_headers(_pdf, _PDF_COLS, _col_widths, _hdr_h)
        _page_bottom = _pdf.h - 16  # footer sits at -14

        for _row_idx, (_, _data_row) in enumerate(df.iterrows()):
            _row_vals = [str(_data_row.get(c, "") or "") for c in _PDF_COLS]
            _max_lines = max(
                len(_pdf.multi_cell(w, _line_h, v, dry_run=True, output="LINES"))
                for v, w in zip(_row_vals, _col_widths)
            )
            _row_h = max(_max_lines * _line_h, _line_h)

            # Page break — close outer borders, add continuation header
            if _pdf.get_y() + _row_h > _page_bottom:
                _pdf.set_draw_color(*_PBR)
                _pdf.set_line_width(0.3)
                _pdf.line(_pdf.l_margin, _y_table_start,
                          _pdf.l_margin, _pdf.get_y())
                _pdf.line(_pdf.l_margin + _total_w, _y_table_start,
                          _pdf.l_margin + _total_w, _pdf.get_y())
                _pdf.add_page()
                _pdf.set_font("DejaVu", size=8.5)
                _pdf.set_text_color(*_PS)
                _pdf.cell(0, 5,
                          "German Vocabulary Export \u2014 continued",
                          new_x="LMARGIN", new_y="NEXT")
                _pdf.ln(3)
                _y_table_start = _pdf.get_y()
                _pdf_col_headers(_pdf, _PDF_COLS, _col_widths, _hdr_h)
                _page_bottom = _pdf.h - 16

            # Full-height alt-row fill
            _x0 = _pdf.l_margin
            _y0 = _pdf.get_y()
            if _row_idx % 2 == 1:
                _pdf.set_fill_color(*_PA)
                _pdf.rect(_x0, _y0, _total_w, _row_h, "F")

            # Cell text — generous left padding
            _pdf.set_font("DejaVu", size=9)
            _pdf.set_text_color(*_PT)
            for v, w in zip(_row_vals, _col_widths):
                _pdf.multi_cell(
                    w, _line_h, f"   {v}",
                    border=0, fill=False,
                    new_x="RIGHT", new_y="TOP",
                    max_line_height=_line_h,
                    align="L",
                )

            # Faint row separator
            _pdf.set_draw_color(*_PBL)
            _pdf.set_line_width(0.15)
            _pdf.line(_x0, _y0 + _row_h, _x0 + _total_w, _y0 + _row_h)
            _pdf.set_y(_y0 + _row_h)

        # Outer border — soft, restrained
        _pdf.set_draw_color(*_PBR)
        _pdf.set_line_width(0.3)
        _ye = _pdf.get_y()
        _pdf.line(_pdf.l_margin, _y_table_start,
                  _pdf.l_margin, _ye)
        _pdf.line(_pdf.l_margin + _total_w, _y_table_start,
                  _pdf.l_margin + _total_w, _ye)
        _pdf.line(_pdf.l_margin, _ye,
                  _pdf.l_margin + _total_w, _ye)

        pdf_bytes = bytes(_pdf.output())

        # ── Download cards — 3 × 2 grid ───────────────────────────────
        dl_left, dl_right = st.columns(2)

        with dl_left:
            st.markdown("""
<div class="dl-card">
  <div class="dl-title">🃏 Anki import TSV</div>
  <div class="dl-desc">Two-column, headerless TSV — Front (English + notes) and Back (German with article). Ready to import into any Anki deck.</div>
</div>""", unsafe_allow_html=True)
            st.download_button(
                label="Download TSV",
                data=tsv_bytes,
                file_name="anki_vocab_export.tsv",
                mime="text/tab-separated-values",
                use_container_width=True,
                key="dl_tsv",
            )

        with dl_right:
            st.markdown("""
<div class="dl-card">
  <div class="dl-title">📄 Full vocabulary CSV</div>
  <div class="dl-desc">All 5 fields as a UTF-8 CSV — opens cleanly in Excel, Google Sheets, or any data tool. German/Afrikaans characters preserved.</div>
</div>""", unsafe_allow_html=True)
            st.download_button(
                label="Download CSV",
                data=csv_bytes,
                file_name="vocab_export.csv",
                mime="text/csv",
                use_container_width=True,
                key="dl_csv",
            )

        dl_left2, dl_right2 = st.columns(2)

        with dl_left2:
            st.markdown("""
<div class="dl-card">
  <div class="dl-title">📊 Full vocabulary XLSX</div>
  <div class="dl-desc">All 5 fields per word — Deutsch, Artikel, Englisch, Afrikaans, Hinweise. Use as a reference or re-import later.</div>
</div>""", unsafe_allow_html=True)
            st.download_button(
                label="Download XLSX",
                data=xlsx_bytes,
                file_name="vocab_export.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="dl_xlsx",
            )

        with dl_right2:
            st.markdown("""
<div class="dl-card">
  <div class="dl-title">🖹 Full vocabulary Markdown</div>
  <div class="dl-desc">All 5 fields as a GFM pipe table — paste into any Markdown editor, README, or Obsidian note.</div>
</div>""", unsafe_allow_html=True)
            st.download_button(
                label="Download Markdown",
                data=md_bytes,
                file_name="vocab_export.md",
                mime="text/markdown",
                use_container_width=True,
                key="dl_md",
            )

        dl_left3, dl_right3 = st.columns(2)

        with dl_left3:
            st.markdown("""
<div class="dl-card">
  <div class="dl-title">📝 Full vocabulary DOCX</div>
  <div class="dl-desc">All 5 fields as a Word document — formatted table, ready to share or print.</div>
</div>""", unsafe_allow_html=True)
            st.download_button(
                label="Download DOCX",
                data=docx_bytes,
                file_name="vocab_export.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_docx",
            )

        with dl_right3:
            st.markdown("""
<div class="dl-card">
  <div class="dl-title">🖨 Full vocabulary PDF</div>
  <div class="dl-desc">All 5 fields as an A4 landscape PDF — formatted table, ready to print or share.</div>
</div>""", unsafe_allow_html=True)
            st.download_button(
                label="Download PDF",
                data=pdf_bytes,
                file_name="vocab_export.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="dl_pdf",
            )

        # ── DB import ──────────────────────────────────────────────────
        st.markdown('<div style="height:0.75rem"></div>', unsafe_allow_html=True)
        st.markdown(
            '<div style="font-size:0.65rem; font-weight:700; letter-spacing:0.1em; '
            'text-transform:uppercase; color:color-mix(in srgb, var(--text-color) 35%, transparent); margin-bottom:0.5rem;">'
            'Import to practice database</div>',
            unsafe_allow_html=True,
        )

        db_left, db_right = st.columns([3, 2])
        source_label_manual = db_left.text_input(
            "Source label",
            value="vocab_export",
            help="Stored as pipeline:<label> in the DB.",
            placeholder="e.g. session_01",
            key="manual_db_source",
            label_visibility="collapsed",
        )
        import_btn = db_right.button(
            "Import to DB →",
            type="secondary",
            use_container_width=True,
            key="manual_db_import_btn",
        )

        if import_btn:
            _do_db_import(df, source_label_manual)
